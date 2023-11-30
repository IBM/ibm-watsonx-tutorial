from django.http import JsonResponse
import csv
import docx
import os
from WatsonxPrj import settings


from dotenv import load_dotenv
load_dotenv(dotenv_path="../.env")

def _generate_csv_file(docx_files:list, csv_folder:str):
        '''
        Create a CSV file from all CV docx files
        '''
        # Create a CSV file to store the records
        csv_file=f'{settings.BASE_DIR}/{csv_folder}/cv_data.csv'
        with open(csv_file, mode='w+', newline='') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerow(["file_name", "text_content"])  # Write header
            for file_name in docx_files:
                text = _extract_text(file_name)
                file_base_name= os.path.basename(file_name) 
                #print([file_base_name, text])
                writer.writerow([file_base_name, text])
            csvfile.close()
        return csv_file
def _extract_text(doc_file_path:str):
        '''
        Extract text from  .docx file
        '''
        texts=[]
        doc_file_path=f'{settings.BASE_DIR}/{doc_file_path}'
        #print(doc_file_path)
        doc = docx.Document(doc_file_path)
        for para in doc.paragraphs:
            texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return "\n".join(texts)
    
def process_cv_files(request):
    try:
        docx_files = request.data.get('cv_docx_files') 
        cv_csv_folder = request.data.get('cv_csv_folder') 
        cv_wxd_table = request.data.get('cv_wxd_table') 
        cv_wxd_target_table = request.data.get('cv_wxd_target_table') 

        if not docx_files:
            return {"error": "No docx_files provided"}

        csv_file=_generate_csv_file(docx_files, cv_csv_folder)
        
        def _upload_csv_table(csv_file_name:str, table_name:str, target_table_name:str):
            try:
                import requests
                watsonx_data_host=os.environ.get("WATSONX_DATA_HOST")
                watsonx_data_inst_id=os.environ.get("WATSONX_DATA_INST_ID")
                watsonx_data_access_token=os.environ.get("WATSONX_DATA_ACCESS_TOKEN")
                
                print(watsonx_data_host)
                print(watsonx_data_inst_id)
                print(watsonx_data_access_token)
                print(table_name)
                print(target_table_name)
                print(csv_file_name)
                  
                headers = {
                    'accept': 'application/json',
                    # 'Content-Type': 'multipart/form-data',
                    'AuthInstanceId': f'{watsonx_data_inst_id}',
                    'LhInstanceId': f'{watsonx_data_inst_id}',
                    'Authorization': f'Bearer {watsonx_data_access_token}',
                    }
                params = {
                    'engine': 'presto-01'
                }
                data = {
                    'catalog': 'iceberg_data',
                    'schema': 'wx_profiles',
                    'tableName': table_name,
                    'ingestionJobName': 'IngestionJob',
                    'scheduled': '1698245204',
                    'created_by': 'admin',
                    'targetTable': target_table_name,
                    'headers': '[ \
                                {"header":"file_name","originalHeader":"file_name","dataType":"varchar"}, \
                                {"header":"text_content","originalHeader":"text_content","dataType":"varchar"} \
                                ]'
                }
                
                files={'csv':open(csv_file_name,'rb')}
                response = requests.post(
                    f'{watsonx_data_host}/lakehouse/api/v1/{watsonx_data_inst_id}/v2/upload/csv',
                    params=params,
                    headers=headers,
                    data=data,
                    files=files
                )
                if( response.status_code == 200):
                    return {"message": "Data processed and uploaded successfully"}
                else:
                    return {"error": f"Failed to upload CV CSV file in watsonx.data. REASON --> {response._content}"} 
            except Exception as ex:
                print("Error in _upload_csv_table:", str(ex))
                raise ex
        #---End _upload_csv_table
        
        try :
            return _upload_csv_table(csv_file, cv_wxd_table,cv_wxd_target_table)
        except Exception  as ex:
            raise ex
    except Exception as e:
        return {"error": str(e)}
    

if __name__=="__main__": 
    
    cv_docs_files=["../data/cv1.docx", "../data/cv2.docx", "../data/cv3.docx", "../data/cv4.docx", "../data/cv5.docx"]
    _generate_csv_file(cv_docs_files, '../data')
    
